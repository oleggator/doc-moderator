# Author: Kirill Syomin

import re
from docx.enum.text import WD_COLOR_INDEX

def insert_paragraph_with_highlighted_words(paragraph_below, text_to_insert, words_list): # (1)
    inserted_paragraph = paragraph_below.insert_paragraph_before()
    mask = '|'.join(words_list)
    matches = re.finditer(mask, text_to_insert)
    prev_end = 0
    for match in matches:
        start = match.start()
        end = match.end()
        inserted_paragraph.add_run(text_to_insert[prev_end:start])
        highlighted_word = inserted_paragraph.add_run(text_to_insert[start:end])
        highlighted_word.font.highlight_color = WD_COLOR_INDEX.RED
        prev_end = end
    inserted_paragraph.add_run(text_to_insert[prev_end:])

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

# в функцию один передаешь docx.paragraph, str (обычно текст текущего параграфа), 
# и список слов для выделения в тексте

# в результате, над параграфом paragraph_below появится параграф с выделенным текстом
# тогда, параграф, который был передан в paragraph_below нужно удалить с помощью
# функции delete_paragraph()