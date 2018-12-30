from Levenshtein import distance
from docx import Document
from os import path
import argparse

from insert_delete import insert_paragraph_with_highlighted_words, delete_paragraph


def find_words(target_words_file_path, article_path, output_path, ld):
    extension = path.splitext(article_path)[1]
    if extension == '.docx':
        document = Document(article_path)
        article_words = set({})
        for paragraph in document.paragraphs:
            for word in split(paragraph.text):
                article_words.add(word)
    else:
        print('wrong file type: .docx only')
        return

    with open(target_words_file_path, 'r') as file:
        target_words = list(set(split(file.read())))

    found_words = set()
    for article_word in article_words:
        for target_word in target_words:
            dist = distance(target_word, article_word)
            if dist <= ld:
                found_words.add(article_word)

    found_words_list = list(found_words)
    for paragraph in document.paragraphs:
        insert_paragraph_with_highlighted_words(paragraph, paragraph.text, found_words_list)
        delete_paragraph(paragraph)

    document.save(output_path)

    print('\n'.join(found_words_list))


def split(test):
    pun = '!"#$%&\'()*+,–./:;<=>?@[\\]^_`{|}~”'
    translator = str.maketrans(pun, ' ' * len(pun))
    return test.lower().translate(translator).split()


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='find words in document')
    parser.add_argument('-d',
                        required=True,
                        action='store',
                        dest='target_words',
                        help='dictionary (txt)')
    parser.add_argument('-a',
                        required=True,
                        action='store',
                        dest='article',
                        help='article (txt or docx)')
    parser.add_argument('-o',
                        required=True,
                        action='store',
                        dest='output',
                        help='found words output file')
    parser.add_argument('-l',
                        required=True,
                        action='store',
                        dest='distance',
                        type=int,
                        help='editor distance')
    args = parser.parse_args()

    find_words(args.target_words, args.article, args.output, args.distance)
